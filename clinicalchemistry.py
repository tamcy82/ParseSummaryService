# Clinical Chemistry Module
import os

# Path: clinicalchemistry.py
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from nltk.tokenize import word_tokenize
from fuzzywuzzy import fuzz
import re

import blood_tube


# Define Clinical Chemistry Class
class ClinicalChemistryLab:

    # Constructor
    def __init__(self, lst, db_path='.\\LocalLabTestsDB.xlsx'):
        # Define paths
        self.study_paths = ['\\\\ctc-network.intranet\\dfs\\BIOTR\\01 Ongoing Studies\\',
                            '\\\\ctc-network.intranet\\dfs\\BIOTR\\02 Closed Studies']
        # Define local service tracker data
        self.lst = lst
        # Define Clinical Chemistry tests
        self.chem_test = []
        # Load ClinicalChemistryTestsDB
        self.ClinicalChemistryTestsDBPath = db_path
        self.ClinicalChemistryTestsDB = self.load_ec_path_code(self.ClinicalChemistryTestsDBPath)
        # Select only clinical chemistry tests
        self.ClinicalChemistryTestsDB = self.ClinicalChemistryTestsDB.loc[self.ClinicalChemistryTestsDB['lab'] == 'chem']
        # Convert all_name to string
        self.ClinicalChemistryTestsDB['alt_name'] = self.ClinicalChemistryTestsDB['alt_name'].astype(str)
        # Expand alt_name column
        self.ClinicalChemistryTestsDB['alt_name'] = self.ClinicalChemistryTestsDB['alt_name'].str.split(',')
        # Define form template path
        self.test_form_template_path = '.\\FormTemplateChemLab.docx'

    # Load ClinicalChemistryTestsDB into pandas
    def load_chemistry_tests_db(self, path):
        df = pd.read_excel(path, header=0)
        return df

    # Load ECPathCode into pandas
    def load_ec_path_code(self, file_path):
        df = pd.read_excel(file_path, header=0, dtype=str, na_filter=False).fillna("")
        # Parse file
        # Create a new dataframe
        df_new = pd.DataFrame(columns=['lab', 'test', 'alt_name', 'specimen', 'code', 'section_order', 'is_optional',
                                       'remarks'])
        # Iterate through each row
        for index, row in df.iterrows():
            # Get lab
            lab = 'chem'
            # Get test
            test = row['Test']
            # Get alt_name
            alt_name = row['Alternative name']
            # Get specimen
            specimen = row['Specimen']
            # Get code
            code = row['ECPath Code']
            # Get section_order
            section_order = row['Section']
            # Get is_optional
            is_optional = ""
            # Get remarks
            remarks = row['Remark on Form'] + " " + row['Internal Remark']
            # Append to new dataframe
            df_new = df_new.append({'lab': lab, 'test': test, 'alt_name': alt_name, 'specimen': specimen,
                                    'code': code, 'section_order': section_order, 'is_optional': is_optional,
                                    'remarks': remarks}, ignore_index=True)
        return df_new

    # Open a given clinical chemistry form template
    def open_chem_form_template(self, path):
        document = Document(path)
        return document

    # Interpret tests
    def interpret_tests(self, item, description, all_tests):
        matched_tests = []
        # Match pregnancy tests
        preg_tests = self.match_pregnancy_tests(item, description, all_tests)
        if len(preg_tests) > 0:
            matched_tests.extend(preg_tests)
        # Match glucose tests
        glucose_tests = self.match_glucose_tests(item, description, all_tests)
        if len(glucose_tests) > 0:
            matched_tests.extend(glucose_tests)
        # Match bilirubin tests
        bilirubin_tests = self.match_bilirubin_tests(item, description, all_tests)
        if len(bilirubin_tests) > 0:
            matched_tests.extend(bilirubin_tests)
        # Match calcium tests
        calcium_tests = self.match_calcium_tests(item, description, all_tests)
        if len(calcium_tests) > 0:
            matched_tests.extend(calcium_tests)
        # Match remaining tests
        remaining_tests = self.match_tests_from_db(item, description, all_tests)
        if len(remaining_tests) > 0:
            matched_tests.extend(remaining_tests)
        # If none of the above tests are matched
        if len(matched_tests) == 0:
            matched_tests.append(item)
        return matched_tests

    # Match proceeding and postceeding words
    def get_proceeding_postceeding_words(self, word_index, texts):
        if word_index < len(texts) - 1:
            # Get proceeding and postceeding words
            if word_index > 0:
                proceeding_word = texts[word_index - 1]
            else:
                proceeding_word = ''
            if word_index < len(texts) - 1:
                postceeding_word = texts[word_index + 1]
            else:
                postceeding_word = ''
            return proceeding_word, postceeding_word
        return None

    # Render exclusion keyword list fo a test
    def get_exclusion_list_of_test(self, test):
        # Default to exclude urine and urinary name
        exclusion_list = ['urine', 'urinary']
        if ' ' in test:
            return exclusion_list
        elif test.startswith('['):
            return exclusion_list
        else:
            # Find tests in the database that contains the test name
            tests = self.ClinicalChemistryTestsDB.loc[self.ClinicalChemistryTestsDB['test'].str.contains(test)]
            target_tests = []
            # Iterate through each test
            for index, row in tests.iterrows():
                if ' ' not in row['test']:
                    target_tests.append(row['test'])
                else:
                    # Split the test name
                    test_names = row['test'].split(' ')
                    # Iterate through each test name
                    for test_word in test_names:
                        if test_word != test:
                            exclusion_list.append(test_word.lower())
        return exclusion_list

    # match pregnancy tests
    # To detect incoming item and description that may contain both urine and serum tests
    def match_pregnancy_tests(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Match basic name
        if "pregnancy" in item.lower() or "pregnancy" in description:
            # Pregnancy tests found
            # Check if urine pregnancy test is requested
            if "urine" in item.lower() or "urine" in description:
                # Urine pregnancy test found
                # Check if urine pregnancy test is not already in all_tests
                if "Urine Pregnancy Test" not in all_tests:
                    # Urine pregnancy test not found in all_tests
                    # Add urine pregnancy test to matched_tests
                    matched_tests.append("Urine Pregnancy Test")
            # Check if serum pregnancy test is requested
            if "serum" in item.lower() or "serum" in description:
                # Serum pregnancy test found
                # Check if serum pregnancy test is not already in all_tests
                if "Serum Pregnancy Test" not in all_tests:
                    # Serum pregnancy test not found in all_tests
                    # Add serum pregnancy test to matched_tests
                    matched_tests.append("Serum Pregnancy Test")
        return matched_tests

    # match bilirubin tests
    # To detect incoming item and description that may contain both total and direct bilirubin tests
    # Check preceeding and postceeding words to detect
    # Return matched_tests
    def match_bilirubin_tests(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Define flags
        is_direct_bilirubin = False
        is_indirect_bilirubin = False
        is_total_bilirubin = False
        # Tokenize item
        item_tokens = word_tokenize(item)
        # Remove punctuation
        item_tokens = [token for token in item_tokens if token.isalnum()]
        # To lower
        item_tokens = [token.lower() for token in item_tokens]
        # Simple ECPath code detection first
        if "dbil" in item_tokens or "dbil" in description:
            matched_tests.append("Direct Bilirubin")
            is_direct_bilirubin = True
        if "tbil" in item_tokens or "tbil" in description:
            matched_tests.append("Total Bilirubin")
            is_total_bilirubin = True
        # Match basic name
        if "bilirubin" in item:
            # Bilirubin tests found
            # Loop through test names with index
            for bilirubin_item_index, t in enumerate(item_tokens):
                if t == "bilirubin":
                    # Get proceeding and postceeding words
                    proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(bilirubin_item_index, item_tokens)
                    # Check if total bilirubin is requested
                    if not is_total_bilirubin and (postceeding_word == "total" or postceeding_word == "total"):
                        # Total bilirubin found
                        is_total_bilirubin = True
                        # Check if total bilirubin is not already in all_tests
                        if "Total Bilirubin" not in all_tests:
                            # Total bilirubin not found in all_tests
                            # Add total bilirubin to matched_tests
                            matched_tests.append("Total Bilirubin")
                    # Check if in-direct bilirubin is requested
                    elif not is_indirect_bilirubin and (proceeding_word == "in-direct" or postceeding_word == "in-direct"):
                        # Indirect bilirubin found
                        is_indirect_bilirubin = True
                        # Check if indirect bilirubin is not already in all_tests
                        if "In-direct Bilirubin" not in all_tests:
                            # Indirect bilirubin not found in all_tests
                            # Add indirect bilirubin to matched_tests
                            matched_tests.append("In-direct Bilirubin")
                    # Check if direct bilirubin is requested
                    elif not is_direct_bilirubin and (proceeding_word == "direct" or postceeding_word == "direct"):
                        # Direct bilirubin found
                        is_direct_bilirubin = True
                        # Check if direct bilirubin is not already in all_tests
                        if "Direct Bilirubin" not in all_tests:
                            # Direct bilirubin not found in all_tests
                            # Add direct bilirubin to matched_tests
                            matched_tests.append("Direct Bilirubin")
        if "bilirubin" in description:
            # Bilirubin tests found
            # Loop through test names with index
            for bilirubin_description_index, t in enumerate(description):
                if t == "bilirubin":
                    # Get proceeding and postceeding words
                    proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(bilirubin_description_index, description)
                    # Check if total bilirubin is requested
                    if not is_total_bilirubin and (postceeding_word == "total" or postceeding_word == "total"):
                        # Total bilirubin found
                        is_total_bilirubin = True
                        # Check if total bilirubin is not already in all_tests
                        if "Total Bilirubin" not in all_tests:
                            # Total bilirubin not found in all_tests
                            # Add total bilirubin to matched_tests
                            matched_tests.append("Total Bilirubin")
                    # Check if in-direct bilirubin is requested
                    elif not is_indirect_bilirubin and (proceeding_word == "in-direct" or postceeding_word == "in-direct"):
                        # Indirect bilirubin found
                        is_indirect_bilirubin = True
                        # Check if indirect bilirubin is not already in all_tests
                        if "In-direct Bilirubin" not in all_tests:
                            # Indirect bilirubin not found in all_tests
                            # Add indirect bilirubin to matched_tests
                            matched_tests.append("In-direct Bilirubin")
                    # Check if direct bilirubin is requested
                    elif not is_direct_bilirubin and (proceeding_word == "direct" or postceeding_word == "direct"):
                        # Direct bilirubin found
                        is_direct_bilirubin = True
                        # Check if direct bilirubin is not already in all_tests
                        if "Direct Bilirubin" not in all_tests:
                            # Direct bilirubin not found in all_tests
                            # Add direct bilirubin to matched_tests
                            matched_tests.append("Direct Bilirubin")
        # If nothing is found
        if ("bilirubin" in item or "bilirubin" in description) and\
                not is_direct_bilirubin and not is_indirect_bilirubin and not is_total_bilirubin:
            # Add total bilirubin to matched_tests
            matched_tests.append("Bilirubin (UNKNOWN TYPE)")
        return matched_tests

    # match different types of glucose tests
    # To detect incoming item and description that may contain different types of glucose tests
    # Check preceeding and postceeding words to detect fasting and random glucose tests
    # Return matched_tests
    def match_glucose_tests(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Define flags
        fasting_glucose_flag = False
        random_glucose_flag = False
        csf_glucose_flag = False
        # Tokenize item
        item_tokens = word_tokenize(item)
        # Remove punctuation
        item_tokens = [token for token in item_tokens if token.isalnum()]
        # To lower
        item_tokens = [token.lower() for token in item_tokens]
        # Simple ECPath code detection first
        if "fg" in item or "fg" in description:
            matched_tests.append("Glucose (Fasting)")
            fasting_glucose_flag = True
        if "rg" in item or "rg" in description:
            matched_tests.append("Glucose (Random)")
            random_glucose_flag = True
        if "cglu" in item or "cglu" in description:
            matched_tests.append("CSF Glucose")
            csf_glucose_flag = True
        # Match name in item
        if "glucose" in item:
            # Glucose tests found
            # Loop through test names with index
            for glucose_item_index, t in enumerate(item_tokens):
                if t == "glucose":
                    if glucose_item_index > 0:
                        # Get proceeding and postceeding words
                        proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(glucose_item_index, item_tokens)
                        # Check if fasting glucose is requested
                        if not fasting_glucose_flag and (proceeding_word == "fasting" or postceeding_word == "fasting"):
                            # Fasting glucose found
                            fasting_glucose_flag = True
                            # Check if fasting glucose is not already in all_tests
                            if "Glucose (Fasting)" not in all_tests:
                                # Fasting glucose not found in all_tests
                                # Add fasting glucose to matched_tests
                                matched_tests.append("Glucose (Fasting)")
                        # Check if random glucose is requested
                        if not random_glucose_flag and\
                                (proceeding_word == "random" or postceeding_word == "random" or
                                 proceeding_word == "non-fasted" or postceeding_word == "non-fasted"):
                            # Random glucose found
                            random_glucose_flag = True
                            # Check if random glucose is not already in all_tests
                            if "Glucose (Random)" not in all_tests:
                                # Random glucose not found in all_tests
                                # Add random glucose to matched_tests
                                matched_tests.append("Glucose (Random)")
                        # Check if csf glucose is requested
                        if not csf_glucose_flag and (proceeding_word == "csf" or postceeding_word == "csf"):
                            # CSF glucose found
                            csf_glucose_flag = True
                            # Check if csf glucose is not already in all_tests
                            if "CSF Glucose" not in all_tests:
                                # CSF glucose not found in all_tests
                                # Add csf glucose to matched_tests
                                matched_tests.append("CSF Glucose")
        # Match name in description
        if "glucose" in description:
            # Glucose tests found
            # Loop through test names with index
            for glucose_description_index, t in enumerate(description):
                if t == "glucose":
                    if glucose_description_index > 0:
                        # Get proceeding and postceeding words
                        proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(glucose_description_index, description)
                        # Check if fasting glucose is requested
                        if not fasting_glucose_flag and (proceeding_word == "fasting" or postceeding_word == "fasting"):
                            # Fasting glucose found
                            fasting_glucose_flag = True
                            # Check if fasting glucose is not already in all_tests
                            if "Glucose (Fasting)" not in all_tests:
                                # Fasting glucose not found in all_tests
                                # Add fasting glucose to matched_tests
                                matched_tests.append("Glucose (Fasting)")
                        # Check if random glucose is requested
                        if not random_glucose_flag and (proceeding_word == "random" or postceeding_word == "random"):
                            # Random glucose found
                            random_glucose_flag = True
                            # Check if random glucose is not already in all_tests
                            if "Glucose (Random)" not in all_tests:
                                # Random glucose not found in all_tests
                                # Add random glucose to matched_tests
                                matched_tests.append("Glucose (Random)")
                        # Check if csf glucose is requested
                        if not csf_glucose_flag and (proceeding_word == "csf" or postceeding_word == "csf"):
                            # CSF glucose found
                            csf_glucose_flag = True
                            # Check if csf glucose is not already in all_tests
                            if "CSF Glucose" not in all_tests:
                                # CSF glucose not found in all_tests
                                # Add csf glucose to matched_tests
                                matched_tests.append("CSF Glucose")
        # If nothing is found
        if ("glucose" in item or "glucose" in description) and\
                not fasting_glucose_flag and not random_glucose_flag and not csf_glucose_flag:
            # Add glucose unknown type to matched_tests
            matched_tests.append("Glucose (Fasting or Random?)")
        return matched_tests

    # Match different types of calcium tests
    # To detect incoming item and description that may contain different types of calcium tests
    # Check preceeding and postceeding words to detect ionized calcium tests
    # Return matched_tests
    def match_calcium_tests(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Define flags
        ionized_calcium_flag = False
        # Tokenize item
        item_tokens = word_tokenize(item)
        # Remove punctuation
        item_tokens = [token for token in item_tokens if token.isalnum()]
        # To lower
        item_tokens = [token.lower() for token in item_tokens]
        # Simple ECPath code detection first
        if "ca" in item or "ca" in description:
            matched_tests.append("Calcium")
        # Match name in item
        if "calcium" in item:
            # Calcium tests found
            # Loop through test names with index
            for calcium_item_index, t in enumerate(item_tokens):
                if t == "calcium":
                    if calcium_item_index > 0:
                        # Get proceeding and postceeding words
                        proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(calcium_item_index, item_tokens)
                        # Check if ionized calcium is requested
                        if not ionized_calcium_flag and (proceeding_word == "ionized" or postceeding_word == "ionized"):
                            # Ionized calcium found
                            ionized_calcium_flag = True
                            # Check if ionized calcium is not already in all_tests
                            if "Ionized Calcium" not in all_tests:
                                # Ionized calcium not found in all_tests
                                # Add ionized calcium to matched_tests
                                matched_tests.append("Calcium (Ionized)")
                        else:
                            # Ionized calcium not found
                            # Check if calcium is not already in all_tests
                            if "Calcium" not in all_tests:
                                # Calcium not found in all_tests
                                # Add calcium to matched_tests
                                matched_tests.append("Calcium")
        # Match name in description
        if "calcium" in description:
            # Calcium tests found
            # Loop through test names with index
            for calcium_description_index, t in enumerate(description):
                if t == "calcium":
                    if calcium_description_index > 0:
                        # Get proceeding and postceeding words
                        proceeding_word, postceeding_word = self.get_proceeding_postceeding_words(calcium_description_index, description)
                        # Check if ionized calcium is requested
                        if not ionized_calcium_flag and (proceeding_word == "ionized" or postceeding_word == "ionized"):
                            # Ionized calcium found
                            ionized_calcium_flag = True
                            # Check if ionized calcium is not already in all_tests
                            if "Ionized Calcium" not in all_tests:
                                # Ionized calcium not found in all_tests
                                # Add ionized calcium to matched_tests
                                matched_tests.append("Calcium (Ionized)")
                        else:
                            # Ionized calcium not found
                            # Check if calcium is not already in all_tests
                            if "Calcium" not in all_tests:
                                # Calcium not found in all_tests
                                # Add calcium to matched_tests
                                matched_tests.append("Calcium")
        # If nothing is found
        if len(matched_tests) == 0\
                and ("calcium" in item or "calcium" in description):
            # Add calcium unknown type to matched_tests
            matched_tests.append("Calcium (UNKNOWN TYPE)")
        return matched_tests

    # Match urine chemistry tests
    # Return matched_tests
    def match_urine_chemisry(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Define test
        urine_tests = {
            "Urinary Protein" : ["urine protein", "urinary protein"],
            "Urinary Creatinine" : ["urine creatinine", "urinary creatinine"],
            "Urine Toxicology Screening" : ["urine toxicology screening", "urine toxicology"],
        }
        scan_range = 3
        # Match item
        for test_name, test_keywords in urine_tests.items():
            for test_keyword in test_keywords:
                if test_keyword in item:
                    # Test found in item
                    # Check if test is not already in all_tests
                    if test_name not in all_tests:
                        actual_name = test_name
                        # Scan description for keywords: 24, spot
                        for scan_index in range(1, scan_range + 1):
                            if "24" in description[scan_index]:
                                actual_name = "24 Hour " + actual_name
                                break
                            elif "spot" in description[scan_index]:
                                actual_name = "Spot " + actual_name
                                break
                        # Test not found in all_tests
                        # Add test to matched_tests
                        matched_tests.append(test_name)
        # Match description
        for test_name, test_keywords in urine_tests.items():
            for test_keyword in test_keywords:
                if " " in test_keyword:
                    # Split test keyword
                    test_keyword_tokens = test_keyword.split(" ")
                    # Loop through tokens
                    for token_index, token in enumerate(test_keyword_tokens):
                        # Compare token with description sequentially
                        if token in description:
                            # Check if token is not the last token
                            if token_index < len(test_keyword_tokens) - 1:
                                # Check if next token is in description
                                if test_keyword_tokens[token_index + 1] in description:
                                    # Check if test is not already in all_tests
                                    actual_name = test_name
                                    # Scan description for keywords: 24, spot
                                    for scan_index in range(token_index - scan_range, token_index + scan_range + 1):
                                        if scan_index >= 0 and scan_index < len(description):
                                            if description[scan_index] == "24":
                                                actual_name = actual_name + " (24-h Urine)"
                                                break
                                            elif description[scan_index] == "spot":
                                                actual_name = actual_name + " (Spot Urine)"
                                                break
                                    if actual_name not in all_tests:
                                        # Test not found in all_tests
                                        # Add test to matched_tests
                                        matched_tests.append(actual_name)
        return matched_tests

    # Match tests from DB
    # To detect incoming item and description that may contain tests from DB
    # Return matched_tests
    def match_tests_from_db(self, item, description, all_tests):
        # Define matched_tests container
        matched_tests = []
        # Exclusion keywords
        # Skip tests that contain these keywords
        scan_range = 1
        # Loop through tests in ClinicalChemistryTestsDB
        for index, test in self.ClinicalChemistryTestsDB.iterrows():
            exclusion_keywords = self.get_exclusion_list_of_test(test['test'])
            # Skip special groups
            if test['test'].startswith("["):
                continue
            # This matching sequence is important
            # Match test name
            this_name = test['test'].lower()
            # Check item
            if item.lower().startswith(this_name):
                # Does not work with tests that have space in their name as description is tokenized
                # Will handle this case later
                # Add to matched tests if not exists
                Excluded = False
                if ' ' in item:
                    # Split item
                    item_tokens = item.split(' ')
                    # Check the exclusion keywords with range
                    for scan_index in range(1, scan_range + 1):
                        if scan_index < len(item_tokens):
                            if item_tokens[scan_index].lower() in exclusion_keywords:
                                Excluded = True
                                continue
                if not Excluded:
                    if test['test'] not in all_tests:
                        # Test not found in all_tests
                        # Add test to matched_tests
                        matched_tests.append(test['test'])
                    continue
            # Check description
            if this_name in description:
                # Does not work with tests that have space in their name as description is tokenized
                # Will handle this case later
                # Add to matched tests if not exists
                Excluded = False
                # Get index of this_name in description
                index = description.index(this_name)
                # Check exclusion keywords with scan range
                for scan_index in range(index - scan_range, index + scan_range + 1):
                    if scan_index >= 0 and scan_index < len(description):
                        if description[scan_index].lower() in exclusion_keywords:
                            Excluded = True
                            continue
                if not Excluded:
                    if test['test'] not in all_tests:
                        # Test not found in all_tests
                        # Add test to matched_tests
                        matched_tests.append(test['test'])
                    continue
            # Match alternative names
            if len(test['alt_name']) > 0:
                for alt_name in test['alt_name']:
                    if alt_name == "":
                        continue
                    if alt_name.lower() in item.lower() or alt_name.lower() in description:
                        matched_tests.append(test['test'])

                        continue
            # For test names have space
            if " " in this_name:
                # Split test into words
                words = this_name.split(" ")
                # Compare test with description sequentially and exclude tests that contain exclusion keywords
                for word_index, word in enumerate(words):
                    if word in description:
                        # Check if word is not the last word
                        if word_index < len(words) - 1:
                            # Check if next word is in description
                            if words[word_index + 1] in description:
                                # Check exclusion keywords with scan range
                                Excluded = False
                                for scan_index in range(word_index - scan_range, word_index + scan_range + 1):
                                    if scan_index >= 0 and scan_index < len(description):
                                        if description[scan_index] in exclusion_keywords:
                                            Excluded = True
                                            continue
                                # Check if test is not already in all_tests
                                if not Excluded:
                                    if test['test'] not in all_tests:
                                        # Test not found in all_tests
                                        # Add test to matched_tests
                                        matched_tests.append(test['test'])
                                    continue
            # Match test name that is not start with the name
            if this_name in item.lower() and not item.lower().startswith(this_name):
                # Check if test is not already in all_tests
                if test['test'] not in all_tests:
                    # Test not found in all_tests
                    # Add test to matched_tests
                    matched_tests.append(test['test'])
                continue
        return matched_tests

    # To modifier the test list
    def chem_test_modifier(self, t):
        # Define special tests
        special_tests = {
            "[TSH, Free T4]": [["TSH"], ["Free T4", "FT4"]],
            "[Cholesterol, Triglycerides, HDL-C, LDL-C]" : [
                ["Cholesterol"], ["Triglycerides"], ["HDL-C"], ["LDL-C"]
            ],
            "[Sodium, Potassium, Chloride, Urea, Creatinine]": [
                ["Sodium"], ["Potassium"], ["Chloride"], ["Urea"], ["Creatinine"]
            ],
            "[Total Protein, Albumin, Total Bilirubin, ALP, AST, ALT]" : [
                ["Total Protein"], ["Albumin"], ["Total Bilirubin"], ["Alkaline Phosphatase", "ALP"], ["AST"], ["ALT"]
            ],
        }
        # Find all special tests from each dictionary values
        for new_name, special_test in special_tests.items():
            # Check if special test is in t
            all_matched = False
            sub_matched = 0
            for sub_test in special_test:
                for test in sub_test:
                    if test in t:
                        sub_matched += 1
                        break
            if sub_matched == len(special_test):
                all_matched = True
            if all_matched:
                # Remove all tests in special_test from t
                for sub_test in special_test:
                    for test in sub_test:
                        if test in t:
                            t.remove(test)
                # Add new_name to t at the beginning
                t.insert(0, new_name)
        return t

    # Render chemistry lab test groups
    def render_test_group(self, T):
        test_group = []
        T = self.chem_test_modifier(T)
        for test in T:
            # Use multiple method to find test in HaematologyTestsDB
            # Exact match
            # Search for test in HaematologyTestsDB
            search_test = self.ClinicalChemistryTestsDB.loc[self.ClinicalChemistryTestsDB['test'] == test]
            # Check alt_name column if not found
            if len(search_test) == 0:
                # Loop through alt_name column
                for index, row in self.ClinicalChemistryTestsDB.iterrows():
                    # Check if test in alt_name
                    if test in row['alt_name']:
                        search_test = self.ClinicalChemistryTestsDB.loc[index]
                        break
            # If not found, clean the test name and search again
            if len(search_test) == 0:
                # Extract alphanumeric and '-' characters from search_test with regex
                test_clean = re.search(r'[\w\- ]+', test).group(0)
                # Trim trailing and leading whitespace
                test_clean = test_clean.strip()
                # Redo search
                search_test = self.ClinicalChemistryTestsDB.loc[self.ClinicalChemistryTestsDB['test'] == test_clean]
                if len(search_test) == 0:
                    for index, row in self.ClinicalChemistryTestsDB.iterrows():
                        if test_clean in row['alt_name']:
                            search_test = self.ClinicalChemistryTestsDB.loc[index]
                            break
            # Find test with similar name
            if len(search_test) == 0:
                for index, row in self.ClinicalChemistryTestsDB.iterrows():
                    if len(test) > 5 and fuzz.token_sort_ratio(test, row['test']) > 80:
                        search_test = self.ClinicalChemistryTestsDB.loc[index]
                        break
            # if found
            if len(search_test) > 0:
                this_group = search_test.iloc[0, 5]
                matched_test = False
                # if thisGroup is not empty
                if this_group == "General Chemistry":
                    for index, row in enumerate(test_group):
                        if row["test_group"] == "General Chemistry":
                            matched_test = True
                            row['Tests'].append(search_test.iloc[0])
                            break
                    if not matched_test:
                        test_group.append({'test_group': "General Chemistry", 'Tests': [search_test.iloc[0]]})
                elif this_group == "Lipid":
                    for index, row in enumerate(test_group):
                        if row["test_group"] == "Lipid":
                            matched_test = True
                            row['Tests'].append(search_test.iloc[0])
                            break
                    if not matched_test:
                        test_group.append({'test_group': "General Chemistry", 'Tests': [search_test.iloc[0]]})
                elif this_group == "Hormone-THT":
                    for index, row in enumerate(test_group):
                        if row["test_group"] == "Hormone-THT":
                            matched_test = True
                            row['Tests'].append(search_test.iloc[0])
                            break
                    if not matched_test:
                        test_group.append({'test_group': "Hormone-THT", 'Tests': [search_test.iloc[0]]})
                else:
                    test_group.append({'test_group': None, 'Tests': [search_test.iloc[0]]})
            else:
                test_group.append({'test_group': None, 'Tests': [
                    {'lab': 'chem', 'test': test, 'alt_name': "", 'specimen': "", 'code': "", 'section_order': None,
                     'is_optional': False, 'remarks': ''}]})
        return test_group

    # Render chemistry test form
    def render_chemistry_test_form(self, site, test_groups):
        UseExportPath = True
        # Check form template path
        if self.test_form_template_path == "":
            print("Chemistry form template path is empty")
            return False
        # Check file exists
        if not os.path.exists(self.test_form_template_path):
            print("Chemistry form template path is not valid")
            return False
        ChemForm = self.open_chem_form_template(self.test_form_template_path)
        # Find site from lst
        site_info = self.lst.loc[self.lst['CTC No.'] == site]
        # Extract digits of site
        CtcNoDigit = ''
        # Extract first digits
        for c in site:
            if c.isdigit():
                CtcNoDigit += c
            else:
                break
        # Fill study info
        if len(ChemForm.tables[0].rows[2].cells[1].paragraphs) > 1:
            SiteCell2 = ChemForm.tables[0].rows[2].cells[1].paragraphs[1]
        else:
            SiteCell2 = ChemForm.tables[0].rows[2].cells[1].add_paragraph()
        SiteCellRun = SiteCell2.add_run('CTC' + CtcNoDigit)
        SiteCellRun.bold = True
        SiteCellRun.font.size = Pt(14)
        if len(site_info) > 0:
            if isinstance(site_info.iloc[0, 30], str):
                # Para0 = FormDoc.tables[0].rows[3].cells[1].paragraphs[0]
                # Para0.add_run('Rept Locn').font.size = Pt(9)
                if len(ChemForm.tables[0].rows[3].cells[1].paragraphs) > 1:
                    Para1 = ChemForm.tables[0].rows[3].cells[1].paragraphs[1]
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.italic = True
                    Para1Run.font.size = Pt(14)
                else:
                    Para1 = ChemForm.tables[0].rows[3].cells[1].add_paragraph()
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.italic = True
                    Para1Run.font.size = Pt(14)
            if isinstance(site_info.iloc[0, 1], str):
                ChemForm.tables[1].rows[0].cells[0].text = ''
                Prot1Run = ChemForm.tables[1].rows[0].cells[0].paragraphs[0].add_run('Protocol: ' + site_info.iloc[0, 1])
                Prot1Run.font.size = Pt(10)
                Prot1Run.bold = True
                ChemForm.tables[1].rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(6)
            if isinstance(site_info.iloc[0, 28], str):
                ChemForm.tables[1].rows[3].cells[0].text = 'Contact Person: ' + site_info.iloc[0, 28]
                ChemForm.tables[1].rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            if isinstance(site_info.iloc[0, 29], str):
                ChemForm.tables[1].rows[3].cells[1].text = 'Contact Number: ' + site_info.iloc[0, 29]
                ChemForm.tables[1].rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        # Fill test info
        test_remarks = []
        for tg in test_groups:
            # Content
            row1 = ChemForm.tables[2].add_row()
            row1.cells[0].text = u'\u25a1'
            # Apply style
            for paragraph in row1.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'PMingLiU'
            # Loop through test with index
            content_run = row1.cells[1].paragraphs[0].add_run()
            content_run.font.size = Pt(11)
            CollectionTubes = []
            # Loop through test with index
            RowFilled = False
            for index, test in enumerate(tg['Tests']):
                if len(content_run.text) < 20:
                    content_run.text += test['test'] + ' (' + test['code'].upper() + ')'
                else:
                    content_run = row1.cells[1].add_paragraph().add_run(test['test'] + ' (' + test['code'].upper() + ')')
                    content_run.font.size = Pt(11)
                # Add specimen to collection tubes if not exists
                if isinstance(test['specimen'], str) and test['specimen'] != '':
                    if test['specimen'] not in CollectionTubes:
                        CollectionTubes.append(test['specimen'])
                RowFilled = True
                # Add test remarks if any
                if isinstance(test['remarks'], str) and len(test['remarks']) > 0:
                    test_remarks.append(test['remarks'])
                # Add coma if not last test
                if index < len(tg['Tests']) - 1:
                    content_run.text += ', '
            # Add collection tubes paragraph
            if len(CollectionTubes) > 0:
                # Loop through collection tubes
                for i, tube in enumerate(CollectionTubes):
                    if isinstance(tube, str) and tube != '':
                        para = row1.cells[1].add_paragraph()
                        para.add_run('[').font.size = Pt(11)
                        run1 = para.add_run(tube)
                        run1.font.size = Pt(11)
                        tube_color = blood_tube.get_blood_tube_colour(tube)
                        if tube_color is not None:
                            run1.font.color.rgb = RGBColor(tube_color[0], tube_color[1], tube_color[2])
                        para.add_run(']').font.size = Pt(11)
            # Merge cell 1 with 2
            row1.cells[1].merge(row1.cells[2])
            # Move row
            rowA = ChemForm.tables[2].rows[len(ChemForm.tables[2].rows) - 1]
            rowB = ChemForm.tables[2].rows[len(ChemForm.tables[2].rows) - 2]
            rowA._tr.addnext(rowB._tr)
        # Remove placeholder row 2
        Table = ChemForm.tables[2]._tbl
        RemoveRow = ChemForm.tables[2].rows[2]._tr
        Table.remove(RemoveRow)
        try:
            # Create a file name
            chem_form_file_name = ''
            if len(site_info) > 0:
                if not isinstance(site_info.iloc[0, 1], str):
                    # Convert to string
                    protocol = str(site_info.iloc[0, 1])
                else:
                    protocol = site_info.iloc[0, 1]
                chem_form_file_name = '[AutoGen] ' + site_info.iloc[0, 0] + '_' + site_info.iloc[0, 2] + '_' + protocol + '_ChemForm.docx'
            else:
                chem_form_file_name = '[AutoGen] ' + site + '_ChemForm.docx'
            if not UseExportPath:
                ChemForm.save(chem_form_file_name)
                print("Chemistry Test Form Rendered: " + chem_form_file_name)
            else:
                # find study folder in export path
                StudyFolder = ''
                for ExportPath in self.study_paths:
                    for dirs in os.listdir(ExportPath):
                        if dirs.startswith(CtcNoDigit + '_'):
                            StudyFolder = os.path.join(ExportPath, dirs)
                            # Check sub folder
                            for dirs2 in os.listdir(StudyFolder):
                                if dirs2.startswith('03 '):
                                    StudyFolder = os.path.join(StudyFolder, dirs2)
                                    break
                            break
                        if StudyFolder != '':
                            break
                    if StudyFolder != '':
                        break
                if StudyFolder == '':
                    print('Error: Study folder not found')
                    print('Export to default path')
                    ChemForm.save(chem_form_file_name)
                else:
                    ChemForm.save(os.path.join(StudyFolder, chem_form_file_name))
                    print("Chemistry Test Form Rendered: " + StudyFolder + "\\" + chem_form_file_name)
        except:
            print('Error: File is open')
        return True
