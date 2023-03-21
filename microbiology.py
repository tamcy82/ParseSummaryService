# Microbiology lab module
# Path: microbiology.py
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from fuzzywuzzy import fuzz
import re


# Define Microbiology Lab
class MicrobiologyLab:

    # Constructor
    def __init__(self, lst, db_path = '.\\LocalLabTestsDB.xlsx', rr_path = '.\\MB_RI_Other Tests_20230317.docx'):
        # Define paths
        self.study_paths = ['\\\\ctc-network.intranet\\dfs\\BIOTR\\01 Ongoing Studies\\',
                            '\\\\ctc-network.intranet\\dfs\\BIOTR\\02 Closed Studies']
        # Define local service tracker data
        self.lst = lst
        # Define microbiology tests
        self.micbio_tests = []
        # Load MicrobiologyTestsDB
        self.MicrobiologyTestsDBPath = db_path
        self.MicrobiologyTestsDB = self.load_microbiology_tests_db(self.MicrobiologyTestsDBPath)
        # Select only microbiology tests
        self.MicrobiologyTestsDB = self.MicrobiologyTestsDB.loc[self.MicrobiologyTestsDB['lab'] == 'micbio']
        # Convert all_name to string
        self.MicrobiologyTestsDB['alt_name'] = self.MicrobiologyTestsDB['alt_name'].astype(str)
        # Expand alt_name column
        self.MicrobiologyTestsDB['alt_name'] = self.MicrobiologyTestsDB['alt_name'].str.split(',')
        # Define HCV serology tests
        self.HCVSerologyTests = ["HCV DNA", "HCV Ab", "HCV RNA", "HCV viral load", "Anti-HCV"]
        # Define Reference Range template path
        self.MicrobiologyRRPath = rr_path
        # Define test form template
        self.test_form_template_path = ".\\FormTemplateMicbioLab.docx"
        return

    # Load MicrobiologyTestsDB into pandas
    def load_microbiology_tests_db(self, path):
        df = pd.read_excel(path, header=0)
        return df

    # Open a given microbiology form template
    def open_micro_form_template(self, path):
        document = Document(path)
        return document

    # Interpret tests
    def interpret_tests(self, item, description, all_tests):
        matched_tests = []
        HCVTest = self.match_hcv_serology(item, description, all_tests)
        # if HCVTest is not empty
        if len(HCVTest) > 0:
            # Add all tests to MicrobiologyTests if not exists
            for test in HCVTest:
                matched_tests.append(test)
            return matched_tests
        HTest = self.match_hiv_serology(item, description, all_tests)
        # if HTest is not empty
        if len(HTest) > 0:
            # Add all tests to MicrobiologyTests if not exists
            for test in HTest:
                matched_tests.append(test)
            return matched_tests
        MTest = self.match_microbiology_tests(item, description, all_tests)
        # if MTest is not empty
        if len(MTest) > 0:
            # Add all tests to MicrobiologyTests if not exists
            for test in MTest:
                matched_tests.append(test)
            return matched_tests
        # No test matched
        if len(MTest) == 0:
            item_name = item.strip()
            matched_tests.append(item_name)
        return matched_tests

    # Match case-insensitive microbiology tests from a given list of string
    # Return a list of matched tests
    def match_microbiology_tests(self, item, description, all_tests):
        matched_tests = []
        # Exact match microbiology tests
        for index, T in self.MicrobiologyTestsDB.iterrows():
            # Check item
            if T['test'].lower() in item.lower():
                matched_tests.append(T['test'])
                continue
            # Check description
            # if T has space
            if " " in T['test']:
                # Split T into words
                words = T['test'].split(" ")
                # Compare T with description sequentially
                for i in range(len(words)):
                    # If T is not in description, break
                    if words[i].lower() not in description:
                        break
                    # If T is in description, and it is the last word of T
                    if words[i].lower() in description and i == len(words) - 1:
                        matched_tests.append(T['test'])
            else:
                if T['test'].lower() in description:
                    matched_tests.append(T['test'])
            # Match alternative names
            for alt_name in T['alt_name']:
                if alt_name.lower() in item.lower() or alt_name.lower() in description:
                    matched_tests.append(T['test'])
        return matched_tests

    # Match case-insensitive HCV serology tests from a given list of string
    def match_hcv_serology(self, item, description, all_tests):
        matched_tests = []
        if ("hcv" in item.lower() or "hcv" in description) and \
                ("serology" in item.lower() or "serology" in description):
            # Exact match HCV serology tests
            for index, T in enumerate(self.HCVSerologyTests):
                if T.lower() in item.lower():
                    matched_tests.append(T)
                    continue
                # Check if T has space
                if " " in T:
                    # Split T into words
                    words = T.split(" ")
                    # Compare T with description sequentially
                    for i in range(len(words)):
                        # If T is not in description, break
                        if words[i].lower() not in description:
                            break
                        # If T is in description, and it is the last word of T
                        if words[i].lower() in description and i == len(words) - 1:
                            matched_tests.append(T)
                else:
                    if T.lower() in description:
                        matched_tests.append(T)
        return matched_tests

    # Match case-insensitive HIV serology tests from a given list of string
    def match_hiv_serology(self, item, description, all_tests):
        if 'hiv' in item.lower() and 'serology' in item.lower():
            return ["HIV 1/2 Ab & HIV Antigen"]
        return []

    # Render a reference range request document
    def render_reference_range_request(self, ctcno, tests, file_name):
        # Open a reference range request template
        document = self.open_micro_form_template(self.MicrobiologyRRPath)
        if len(tests) > 0:
            # Clone the test list
            remaining_tests = tests.copy()
            # Set reference table
            RTable = document.tables[0]
            # Loop through each row of the table in reverse order
            for i in range(len(RTable.rows) - 1, 0, -1):
                # Get the test name
                test = RTable.cell(i, 0).text.strip()
                # If test is not in tests, delete the row
                if test not in tests:
                    RemoveRow = RTable.rows[i]._tr
                    RTable._tbl.remove(RemoveRow)
                else:
                    remaining_tests.remove(test)
            # If there are remaining tests, add them to the table
            if len(remaining_tests) > 0:
                for test in remaining_tests:
                    RTable.add_row()
                    RTable.cell(len(RTable.rows) - 1, 0).text = test
                    # Set cell spacing
                    RTable.cell(len(RTable.rows) - 1, 0).paragraphs[0].paragraph_format.space_before = Pt(6)
                    RTable.cell(len(RTable.rows) - 1, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
                    RTable.cell(len(RTable.rows) - 1, 1).paragraphs[0].paragraph_format.space_before = Pt(6)
                    RTable.cell(len(RTable.rows) - 1, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
                    RTable.cell(len(RTable.rows) - 1, 2).paragraphs[0].paragraph_format.space_before = Pt(6)
                    RTable.cell(len(RTable.rows) - 1, 2).paragraphs[0].paragraph_format.space_after = Pt(6)
            # Set footer
            footer = document.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Reference No.: CTC" + ctcno;
            for r in footer_para.runs:
                r.font.italic = True
            # Save the document
            document.save(file_name)
            return True
        else:
            return False

    # Render microbiology lab test groups
    def render_test_group(self, T):
        TestGroup = []
        for test in T:
            # Use multiple method to find test in HaematologyTestsDB
            # Exact match
            # Search for test in HaematologyTestsDB
            SearchTest = self.MicrobiologyTestsDB.loc[self.MicrobiologyTestsDB['test'] == test]
            # Check alt_name column if not found
            if len(SearchTest) == 0:
                # Loop through alt_name column
                for index, row in self.MicrobiologyTestsDB.iterrows():
                    # Check if test in alt_name
                    if test in row['alt_name']:
                        SearchTest = self.MicrobiologyTestsDB.loc[index]
                        break
            # If not found, clean the test name and search again
            if len(SearchTest) == 0:
                # Extract alphanumeric and '-' characters from SearchTest with regex
                test_clean = re.search(r'[\w\- ]+', test).group(0)
                # Trim trailing and leading whitespace
                test_clean = test_clean.strip()
                # Redo search
                SearchTest = self.MicrobiologyTestsDB.loc[self.MicrobiologyTestsDB['test'] == test_clean]
                if len(SearchTest) == 0:
                    for index, row in self.MicrobiologyTestsDB.iterrows():
                        if test_clean in row['alt_name']:
                            SearchTest = self.MicrobiologyTestsDB.loc[index]
                            break
            # Find test with similar name
            if len(SearchTest) == 0:
                for index, row in self.MicrobiologyTestsDB.iterrows():
                    if len(test) > 5 and fuzz.token_sort_ratio(test, row['test']) > 80:
                        SearchTest = self.MicrobiologyTestsDB.loc[index]
                        break
            # if found
            if len(SearchTest) > 0:
                thisGroup = SearchTest.iloc[0, 3]
                MatchedTest = False
                # if thisGroup is not empty
                if thisGroup is not None:
                    for index, row in enumerate(TestGroup):
                        if row["TestGroup"] == thisGroup:
                            MatchedTest = True
                            row['Tests'].append(SearchTest.iloc[0])
                            break
                    if not MatchedTest:
                        TestGroup.append({'TestGroup': thisGroup, 'Tests': [SearchTest.iloc[0]]})
                else:
                    TestGroup.append({'TestGroup': None, 'Tests': [SearchTest.iloc[0]]})
            else:
                TestGroup.append({'TestGroup': None, 'Tests': [
                    {'test': test, 'description': None, 'specimen': "", 'code': "", 'alt_name': "",
                     'interpretation': None, 'is_optional': False, 'remarks': ''}]})
        return TestGroup

    # Render a microbiology lab test form
    def render_test_form(self, site, test_groups):
        UseExportPath = True
        # Check form template path
        if self.test_form_template_path == "":
            return False
        # Check file exists
        if not os.path.exists(self.test_form_template_path):
            return False
        MicbioForm = self.open_micro_form_template(self.test_form_template_path)
        # Find site from lst
        site_info = self.lst.loc[self.lst['CTC No.'] == site]
        # Extract digits of site
        CtcNoDigit = ''
        # Extract first digits
        for c in site:
            if c.isdigit():
                CtcNoDigit += c
            else:
                break
        # Fill study info
        if len(MicbioForm.tables[0].rows[2].cells[1].paragraphs) > 1:
            SiteCell2 = MicbioForm.tables[0].rows[2].cells[1].paragraphs[1]
        else:
            SiteCell2 = MicbioForm.tables[0].rows[2].cells[1].add_paragraph()
        SiteCellRun = SiteCell2.add_run('CTC' + CtcNoDigit)
        SiteCellRun.bold = True
        SiteCellRun.font.size = Pt(14)
        if len(site_info) > 0:
            if isinstance(site_info.iloc[0, 30], str):
                # Para0 = FormDoc.tables[0].rows[3].cells[1].paragraphs[0]
                # Para0.add_run('Rept Locn').font.size = Pt(9)
                if len(MicbioForm.tables[0].rows[3].cells[1].paragraphs) > 1:
                    Para1 = MicbioForm.tables[0].rows[3].cells[1].paragraphs[1]
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.italic = True
                    Para1Run.font.size = Pt(14)
                else:
                    Para1 = MicbioForm.tables[0].rows[3].cells[1].add_paragraph()
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.italic = True
                    Para1Run.font.size = Pt(14)
            if isinstance(site_info.iloc[0, 1], str):
                MicbioForm.tables[1].rows[0].cells[0].text = ''
                Prot1Run = MicbioForm.tables[1].rows[0].cells[0].paragraphs[0].add_run('Protocol: ' + site_info.iloc[0, 1])
                Prot1Run.font.size = Pt(10)
                Prot1Run.bold = True
                MicbioForm.tables[1].rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(6)
            if isinstance(site_info.iloc[0, 28], str):
                MicbioForm.tables[1].rows[3].cells[0].text = 'Contact Person: ' + site_info.iloc[0, 28]
                MicbioForm.tables[1].rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            if isinstance(site_info.iloc[0, 29], str):
                MicbioForm.tables[1].rows[3].cells[1].text = 'Contact Number: ' + site.iloc[0, 29]
                MicbioForm.tables[1].rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        # Content
        # Get content row
        row1 = MicbioForm.tables[2].rows[1]
        FirstPara = False
        # Loop through test groups
        for tg in test_groups:
            # Print all test group
            # print(tg)
            # Loop through test with index
            CollectionTubes = []
            for index, test in enumerate(tg['Tests']):
                if not FirstPara:
                    FirstPara = True
                    para = row1.cells[0].paragraphs[0]
                else:
                    para = row1.cells[0].add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                if test['code'] != '':
                    run1 = para.add_run(test['code'])
                else:
                    run1 = para.add_run('Unknown')
                run1.font.highlight_color = WD_COLOR_INDEX.PINK
                run1.font.size = Pt(10)
                run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run2 = para.add_run(' ' + test['test'])
                run2.font.size = Pt(10)
                # Add specimen to collection tubes if not exists
                if isinstance(test['specimen'], str) and test['specimen'] != '':
                    if test['specimen'] not in CollectionTubes:
                        CollectionTubes.append(test['specimen'])
            # Add collection tubes paragraph
            if len(CollectionTubes) > 0:
                # Loop through collection tubes
                for i, tube in enumerate(CollectionTubes):
                    if isinstance(tube, str) and tube != '':
                        para = row1.cells[0].add_paragraph('[' + tube + ']')
                        para.runs[0].font.size = Pt(10)
                        para.paragraph_format.space_before = Pt(6)
                        para.paragraph_format.space_after = Pt(6)
        try:
            RRExportFileName = ''
            if (len(site_info) > 0):
                RRExportFileName = '[AutoGen] ' + site_info.iloc[0,0] + '_' + site_info.iloc[0,2] + '_' + site_info.iloc[0,1] + '_MicrobioForm.docx'
            else:
                RRExportFileName = '[AutoGen] ' + site + '_MicrobioForm.docx'
            if not UseExportPath:
                MicbioForm.save(RRExportFileName)
                print("Microbiology Test Form Rendered: " + RRExportFileName)
            else:
                # find study folder in export path
                StudyFolder = ''
                for ExportPath in self.study_paths:
                    for dirs in os.listdir(ExportPath):
                        if dirs.startswith(CtcNoDigit + '_'):
                            StudyFolder = os.path.join(ExportPath, dirs)
                            # Check sub folder
                            for dirs2 in os.listdir(StudyFolder):
                                if dirs2.startswith('03 '):
                                    StudyFolder = os.path.join(StudyFolder, dirs2)
                                    break
                            break
                        if StudyFolder != '':
                            break
                    if StudyFolder != '':
                        break
                if StudyFolder == '':
                    print('Error: Study folder not found')
                    print('Export to default path')
                    MicbioForm.save(RRExportFileName)
                    return
                else:
                    MicbioForm.save(os.path.join(StudyFolder, RRExportFileName))
                    print("Microbiology Test Form Rendered: " + StudyFolder + "\\" + RRExportFileName)
        except:
            print('Error: File is open')
        return True
