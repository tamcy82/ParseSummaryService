# Haematology lab module
# Interpreting haematology lab test
import os

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from fuzzywuzzy import fuzz
import re

import blood_tube


# Path: haematology.py


class HaematologyLab:

    # Define Haematology tests
    def __init__(self, lst, db_path='.\\LocalLabTestsDB.xlsx'):
        # Define paths
        self.study_paths = ['\\\\ctc-network.intranet\\dfs\\BIOTR\\01 Ongoing Studies\\',
                            '\\\\ctc-network.intranet\\dfs\\BIOTR\\02 Closed Studies']
        # Define coagulation tests
        self.CoagulationTests = ["PT", "aPTT", "INR", "Fibrinogen", "D-Dimer"]
        self.TypingTests = ["ABO", "RhD", "Phenotyping", "Antibody Screening", "Direct Antiglobulin Test",
                            "Indirect Antiglobulin Test"]
        # Define local service tracker data
        self.lst = lst
        # Load HaematologyTestsDB
        self.HaematologyTestsDBPath = db_path
        self.HaematologyTestsDB = self.load_haematology_tests_db(self.HaematologyTestsDBPath)
        # Select only haematology tests
        self.HaematologyTestsDB = self.HaematologyTestsDB.loc[self.HaematologyTestsDB['lab'] == 'haema']
        # Convert all_name to string
        self.HaematologyTestsDB['alt_name'] = self.HaematologyTestsDB['alt_name'].astype(str)
        # Expand alt_name column
        self.HaematologyTestsDB['alt_name'] = self.HaematologyTestsDB['alt_name'].str.split(',')
        # Test form path
        self.test_form_template_path = '.\\FormTemplateHaemaLab.docx'

    # Load HaematologyTestsDB into pandas
    def load_haematology_tests_db(self, path):
        df = pd.read_excel(path, header=0)
        return df

    # Open form word file
    def open_haema_form_template(self, path):
        document = Document(path)
        return document

    # Interpret tests
    def interpret_tests(self, item, description, all_tests):
        matched_tests = []
        # Check special tests
        CTest = self.match_coagulation_tests(item, description, all_tests)
        # if CTest is not empty
        if len(CTest) > 0:
            # Add all tests to HaematologyTests if not exists
            for test in CTest:
                matched_tests.append(test)
        # Blood Film Examination
        BloodFilms = self.match_blood_film_examination(item, description, all_tests)
        if len(BloodFilms) > 0:
            # Add all tests to HaematologyTests if not exists
            for test in BloodFilms:
                matched_tests.append(test)
        # Complete Blood Picture
        CBC = self.match_complete_blood_picture(item, description, all_tests)
        if len(CBC) > 0:
            # Add all tests to HaematologyTests if not exists
            for test in CBC:
                matched_tests.append(test)
        # Typing
        BTyping = self.match_blood_typing(item, description, all_tests)
        if len(BTyping) > 0:
            # Add all tests to HaematologyTests if not exists
            for test in BTyping:
                matched_tests.append(test)
        # None matched
        if len(CTest) == 0 and len(BloodFilms) == 0 and len(CBC) == 0 and len(BTyping) == 0:
            # Add test to HaematologyTests if not exists
            item_name = item.strip()
            matched_tests.append(item_name)
        return matched_tests

    # Match case-insensitive coagulation tests from a given list of string
    # Return a list of matched tests
    def match_coagulation_tests(self, item, description, all_tests):
        matched_tests = []
        # Exact match coagulation tests
        for index, T in enumerate(self.CoagulationTests):
            if T.lower() in item.lower() or T.lower() in description:
                matched_tests.append(self.CoagulationTests[index])
        return matched_tests

    # Match blood film examination
    def match_blood_film_examination(self, item, description, all_tests):
        if 'examination of blood film' not in item.lower():
            return []
        matched_tests = ["Blood Film Examination"]
        # Define flag for blast and morphology
        additional_notes = []
        # Check if string 'blast' in description, case-insensitive with index
        for i in range(len(description)):
            if 'blast' in description[i].lower():
                # Add to additional note if not exists
                if "Blasts" not in additional_notes:
                    additional_notes.append("Blasts")
                    continue
            if 'morphology' in description[i].lower():
                # Add to additional note if not exists
                if "Morphology" not in additional_notes:
                    additional_notes.append("General Morphology")
                    continue
            if 'nucleated' in description[i].lower():
                # Add to additional note if not exists
                if "Nucleated" not in additional_notes:
                    additional_notes.append("Nucleated RBC")
                    continue
        if len(additional_notes) > 0:
            additional_notes.sort()
            matched_tests[0] += " (include"
            # Loop through additional notes with index
            for i in range(len(additional_notes)):
                matched_tests[0] += " " + additional_notes[i]
                # if last two items, add 'and'
                if i == len(additional_notes) - 2:
                    matched_tests[0] += " and"
                # else add comma
                elif i < len(additional_notes) - 1:
                    matched_tests[0] += ","
            matched_tests[0] += " evaluation)"
        return matched_tests

    # Mach complete blood picture
    def match_complete_blood_picture(self, item, description, all_tests):
        if 'haematology' in item.lower() or 'count' in description:
            matched_tests = ["Complete Blood Picture"]
            if 'reticulocyte' in item.lower() or 'reticulocyte' in description:
                matched_tests.append("Reticulocytes")
            if 'fibrinogen' in item.lower() or 'fibrinogen' in description:
                matched_tests.append("Fibrinogen")
            return matched_tests
        elif 'reticulocyte' in item.lower() and "Complete Blood Picture" in all_tests:
            return ["Reticulocytes"]
        else:
            return []

    # Match blood typing
    def match_blood_typing(self, item, description, all_tests):
        matched_tests = []
        # Exact match typing tests
        for index, T in enumerate(self.TypingTests):
            if T.lower() in item.lower() or T.lower() in description:
                matched_tests.append(self.TypingTests[index])
        return matched_tests

    # Match Trephine Biopsy IHC Reporting
    def match_trephine_biopsy(self, description):
        return ["Trephine Biopsy IHC Reporting"]

    # Match Bone Marrow Aspirate Reporting
    def match_bone_marrow_aspirate(self, description):
        return ["Bone Marrow Aspirate Reporting"]

    # Match Bone Marrow Smear Preparation
    def match_bone_marrow_smear(self, description):
        return ["Bone Marrow Smear Preparation"]

    # Match Cytogenetics
    def match_cytogenetics(self, description):
        return ["Cytogenetics"]

    # Match Blood Phenotyping
    def match_blood_phenotyping(self, description):
        return ["Blood Phenotyping"]

    # Create a list of Haematology tests from a text string
    def render_haematology_test_group(self, T):
        TestGroup = []
        for test in T:
            # Use multiple method to find test in HaematologyTestsDB
            # Exact match
            # Search for test in HaematologyTestsDB
            SearchTest = self.HaematologyTestsDB.loc[self.HaematologyTestsDB['test'] == test]
            # Check alt_name column if not found
            if len(SearchTest) == 0:
                # Loop through alt_name column
                for index, row in self.HaematologyTestsDB.iterrows():
                    # Check if test in alt_name
                    if test in row['alt_name']:
                        SearchTest = self.HaematologyTestsDB.loc[index]
                        break
            # If not found, clean the test name and search again
            if len(SearchTest) == 0:
                # Extract alphanumeric and '-' characters from SearchTest with regex
                test_clean = re.search(r'[\w\- ]+', test).group(0)
                # Trim trailing and leading whitespace
                test_clean = test_clean.strip()
                # Redo search
                SearchTest = self.HaematologyTestsDB.loc[self.HaematologyTestsDB['test'] == test_clean]
                if len(SearchTest) == 0:
                    for index, row in self.HaematologyTestsDB.iterrows():
                        if test_clean in row['alt_name']:
                            SearchTest = self.HaematologyTestsDB.loc[index]
                            break
            # Find test with similar name
            if len(SearchTest) == 0:
                for index, row in self.HaematologyTestsDB.iterrows():
                    if len(test) > 5 and fuzz.token_sort_ratio(test, row['test']) > 80:
                        SearchTest = self.HaematologyTestsDB.loc[index]
                        break
            # if found
            if len(SearchTest) > 0:
                this_group = SearchTest.iloc[0, 3]
                MatchedTest = False
                # if this_group is not empty
                if this_group is not None:
                    for index, row in enumerate(TestGroup):
                        if row["TestGroup"] == this_group:
                            MatchedTest = True
                            row['Tests'].append(SearchTest.iloc[0])
                            break
                    if not MatchedTest:
                        TestGroup.append({'TestGroup': this_group, 'Tests': [SearchTest.iloc[0]]})
                else:
                    TestGroup.append({'TestGroup': None, 'Tests': [SearchTest.iloc[0]]})
            else:
                TestGroup.append({'TestGroup': None, 'Tests': [
                    {'lab': 'haema', 'test': test, 'alt_name': "", 'specimen': "", 'code': "", 'section_order': None,
                     'is_optional': False, 'remarks': ''}]})
        return TestGroup

    # Render haematology test form
    def render_haematology_test_form(self, site, test_groups):
        UseExportPath = True
        # Check form template path
        if self.test_form_template_path == "":
            print("Chemistry form template path is empty")
            return False
        # Check file exists
        if not os.path.exists(self.test_form_template_path):
            print("Chemistry form template path is not valid")
            return False
        HaemaForm = self.open_haema_form_template(self.test_form_template_path)
        # Find site from lst
        site_info = self.lst.loc[self.lst['CTC No.'] == site]
        # Extract digits of site
        CtcNoDigit = ''
        # Extract first digits
        for c in site:
            if c.isdigit():
                CtcNoDigit += c
            else:
                break
        # Fill study info
        if len(HaemaForm.tables[0].rows[2].cells[1].paragraphs) > 1:
            SiteCell2 = HaemaForm.tables[0].rows[2].cells[1].paragraphs[1]
        else:
            SiteCell2 = HaemaForm.tables[0].rows[2].cells[1].add_paragraph()
        SiteCell2.add_run('CTC' + CtcNoDigit).bold = True
        # Select row by key in LST
        site_info = self.lst.loc[self.lst['CTC No.'] == site]
        if len(site_info) > 0:
            if isinstance(site_info.iloc[0, 30], str):
                # Para0 = FormDoc.tables[0].rows[3].cells[1].paragraphs[0]
                # Para0.add_run('Rept Locn').font.size = Pt(9)
                if len(HaemaForm.tables[0].rows[3].cells[1].paragraphs) > 1:
                    Para1 = HaemaForm.tables[0].rows[3].cells[1].paragraphs[1]
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.font.size = Pt(12)
                else:
                    Para1 = HaemaForm.tables[0].rows[3].cells[1].add_paragraph()
                    Para1Run = Para1.add_run(site_info.iloc[0, 30])
                    Para1Run.bold = True
                    Para1Run.font.size = Pt(12)
            if isinstance(site_info.iloc[0, 1], str):
                HaemaForm.tables[1].rows[0].cells[0].text = ''
                Prot1Run = HaemaForm.tables[1].rows[0].cells[0].paragraphs[0].add_run('Protocol: ' + site_info.iloc[0, 1])
                Prot1Run.font.size = Pt(10)
                Prot1Run.bold = True
                HaemaForm.tables[1].rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(6)
            if isinstance(site_info.iloc[0, 28], str):
                HaemaForm.tables[1].rows[3].cells[0].text = 'Contact Person: ' + site_info.iloc[0, 28]
                HaemaForm.tables[1].rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            if isinstance(site_info.iloc[0, 29], str):
                HaemaForm.tables[1].rows[3].cells[1].text = 'Contact Number: ' + site_info.iloc[0, 29]
                HaemaForm.tables[1].rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        # Loop through tests
        # Test remarks
        test_remarks = []
        for TG in test_groups:
            # Content
            row1 = HaemaForm.tables[2].add_row()
            row1.cells[0].text = u'\u25a1'
            # Apply style
            for paragraph in row1.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'PMingLiU'
            # Loop through test with index
            Content = ''
            CollectionTubes = []
            OptionalTest = []
            # Check at least one test is not optional
            HasNonOptional = False
            for test in TG['Tests']:
                if not test['is_optional']:
                    HasNonOptional = True
            # Loop through test with index
            RowFilled = False
            for index, test in enumerate(TG['Tests']):
                if HasNonOptional and test['is_optional'] is True:
                    OptionalTest.append(test)
                    continue
                if RowFilled:
                    Content += '\n'
                Content += test['test']
                # Add specimen to collection tubes if not exists
                if isinstance(test['specimen'], str) and test['specimen'] != '':
                    if test['specimen'] not in CollectionTubes:
                        CollectionTubes.append(test['specimen'])
                RowFilled = True
                # Add test remarks if any
                if isinstance(test['remarks'], str) and len(test['remarks']) > 0:
                    test_remarks.append(test['remarks'])
            row1.cells[1].text = Content
            # Add collection tubes paragraph
            if len(CollectionTubes) > 0:
                # Loop through collection tubes
                for i, tube in enumerate(CollectionTubes):
                    if isinstance(tube, str) and tube != '':
                        para = row1.cells[1].add_paragraph()
                        para.add_run('[')
                        run1 = para.add_run(tube)
                        tube_color = blood_tube.get_blood_tube_colour(tube)
                        if tube_color is not None:
                            run1.font.color.rgb = RGBColor(tube_color[0], tube_color[1], tube_color[2])
                        para.add_run(']')
            # Merge cell 1 with 2
            row1.cells[1].merge(row1.cells[2])
            # Move row
            rowA = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 1]
            rowB = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 2]
            rowA._tr.addnext(rowB._tr)
            # Add optional test
            if len(OptionalTest) > 0:
                for i2, t2 in enumerate(OptionalTest):
                    row2 = HaemaForm.tables[2].add_row()
                    row2.cells[1].text = u'\u25a1'
                    # Apply style
                    for paragraph in row2.cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'PMingLiU'
                    row2.cells[2].text = t2['test']
                    # Move row
                    rowA = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 1]
                    rowB = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 2]
                    rowA._tr.addnext(rowB._tr)
                # Merge row
                Cell1 = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 2].cells[0]
                Cell2 = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 3].cells[0]
                Cell1.merge(Cell2)
        # Remove placeholder row 2
        Table = HaemaForm.tables[2]._tbl
        RemoveRow = HaemaForm.tables[2].rows[2]._tr
        Table.remove(RemoveRow)
        # Add test remarks to the last row of the table
        if len(test_remarks) > 0:
            row = HaemaForm.tables[2].rows[len(HaemaForm.tables[2].rows) - 1]
            for R in test_remarks:
                para_remark = row.cells[1].add_paragraph()
                para_run = para_remark.add_run(R)
                para_run.font.italic = True
                para_run.font.bold = True
                para_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                para_run.font.size = Pt(10)
        try:
            RRExportFileName = ''
            if (len(site_info) > 0):
                if not isinstance(site_info.iloc[0, 1], str):
                    # Convert to string
                    protocol = str(site_info.iloc[0, 1])
                else:
                    protocol = site_info.iloc[0, 1]
                RRExportFileName = '[AutoGen] ' + site_info.iloc[0,0] + '_' + site_info.iloc[0,2] + '_' + protocol + '_HemaForm.docx'
            else:
                RRExportFileName = '[AutoGen] ' + site + '_HemaForm.docx'
            if not UseExportPath:
                HaemaForm.save(RRExportFileName)
                print("Haematology Test Form Rendered: " + RRExportFileName)
            else:
                # find study folder in export path
                StudyFolder = ''
                for ExportPath in self.study_paths:
                    for dirs in os.listdir(ExportPath):
                        if dirs.startswith(CtcNoDigit + '_'):
                            StudyFolder = os.path.join(ExportPath, dirs)
                            # Check sub folder
                            for dirs2 in os.listdir(StudyFolder):
                                if dirs2.startswith('03 '):
                                    StudyFolder = os.path.join(StudyFolder, dirs2)
                                    break
                            break
                        if StudyFolder != '':
                            break
                    if StudyFolder != '':
                        break
                if StudyFolder == '':
                    print('Error: Study folder not found')
                    print('Export to default path')
                    HaemaForm.save(RRExportFileName)
                else:
                    HaemaForm.save(os.path.join(StudyFolder, RRExportFileName))
                    print("Haematology Test Form Rendered: " + StudyFolder + "\\" + RRExportFileName)
            return True
        except:
            print('Error: File is open')
            return False